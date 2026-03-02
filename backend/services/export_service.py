"""
简历导出服务 - PDF/Word生成
"""

from typing import Dict, Optional
import os


class ResumeExporter:
    """简历导出器"""
    
    # 模板配置
    TEMPLATES = {
        "default": {
            "name": "通用模板",
            "style": "minimal",
            "ats_friendly": True
        },
        "modern": {
            "name": "现代模板",
            "style": "modern",
            "ats_friendly": True
        },
        "professional": {
            "name": "专业模板",
            "style": "classic",
            "ats_friendly": True
        },
        "creative": {
            "name": "创意模板",
            "style": "creative",
            "ats_friendly": False
        }
    }
    
    def __init__(self):
        self.output_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
            "public", "resumes"
        )
        os.makedirs(self.output_dir, exist_ok=True)
    
    def export_pdf(self, resume_data: Dict, template_id: str = "default") -> str:
        """
        导出PDF简历
        
        Args:
            resume_data: 简历数据
            template_id: 模板ID
        
        Returns:
            PDF文件URL
        """
        # 方案1: 尝试使用 weasyprint（需要安装）
        try:
            from weasyprint import HTML
            return self._export_pdf_weasyprint(resume_data, template_id)
        except ImportError:
            pass
        
        # 方案2: 尝试使用 pdfkit（需要安装 wkhtmltopdf）
        try:
            import pdfkit
            return self._export_pdf_pdfkit(resume_data, template_id)
        except ImportError:
            pass
        
        # 方案3: 降级 - 生成HTML（用户可打印为PDF）
        return self._export_pdf_html(resume_data, template_id)
    
    def _export_pdf_weasyprint(self, resume_data: Dict, template_id: str) -> str:
        """使用 WeasyPrint 生成 PDF"""
        html_content = self._generate_html(resume_data, template_id)
        
        # 生成文件名
        filename = f"resume_{resume_data.get('id', 'draft')}.pdf"
        filepath = os.path.join(self.output_dir, filename)
        
        # 生成PDF
        HTML(string=html_content).write_pdf(filepath)
        
        return f"/resumes/{filename}"
    
    def _export_pdf_pdfkit(self, resume_data: Dict, template_id: str) -> str:
        """使用 pdfkit 生成 PDF"""
        html_content = self._generate_html(resume_data, template_id)
        
        filename = f"resume_{resume_data.get('id', 'draft')}.pdf"
        filepath = os.path.join(self.output_dir, filename)
        
        # 配置 pdfkit（需要系统安装 wkhtmltopdf）
        options = {
            'page-size': 'A4',
            'margin-top': '0.5in',
            'margin-right': '0.5in',
            'margin-bottom': '0.5in',
            'margin-left': '0.5in',
            'encoding': 'utf-8'
        }
        
        pdfkit.from_string(html_content, filepath, options=options)
        
        return f"/resumes/{filename}"
    
    def _export_pdf_html(self, resume_data: Dict, template_id: str) -> str:
        """生成可打印的HTML（降级方案）"""
        html_content = self._generate_html(resume_data, template_id)
        
        filename = f"resume_{resume_data.get('id', 'draft')}.html"
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return f"/resumes/{filename}"
    
    def export_word(self, resume_data: Dict, template_id: str = "default") -> str:
        """
        导出 Word 简历
        
        Args:
            resume_data: 简历数据
            template_id: 模板ID
        
        Returns:
            Word文件URL
        """
        # 尝试使用 python-docx
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            return self._export_docx(resume_data, template_id)
        except ImportError:
            pass
        
        # 降级：返回提示信息
        return "/resumes/comingsoon.docx"
    
    def _export_docx(self, resume_data: Dict, template_id: str) -> str:
        """使用 python-docx 生成 Word"""
        doc = Document()
        
        # 标题
        header = resume_data.get("header", {})
        title = doc.add_heading(header.get("name", "简历"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 联系信息
        contact_info = f"{header.get('phone', '')} | {header.get('email', '')} | {header.get('location', '')}"
        p = doc.add_paragraph(contact_info)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 个人简介
        if resume_data.get("summary"):
            doc.add_heading("个人简介", level=1)
            doc.add_paragraph(resume_data.get("summary"))
        
        # 工作经历
        experience = resume_data.get("experience", [])
        if experience:
            doc.add_heading("工作经历", level=1)
            for exp in experience:
                company = exp.get("company_name", "")
                position = exp.get("position", "")
                duration = f"{exp.get('start_date', '')} - {exp.get('end_date', '至今')}"
                
                p = doc.add_paragraph()
                p.add_run(f"{position} - {company}").bold = True
                p.add_run(f" ({duration})")
                
                for ach in exp.get("achievements", []):
                    doc.add_paragraph(ach, style='List Bullet')
        
        # 教育背景
        education = resume_data.get("education", [])
        if education:
            doc.add_heading("教育背景", level=1)
            for edu in education:
                school = edu.get("school_name", "")
                major = edu.get("major", "")
                degree = edu.get("degree", "")
                duration = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}"
                
                p = doc.add_paragraph()
                p.add_run(f"{school} - {major} ({degree})").bold = True
                p.add_run(f" | {duration}")
        
        # 技能
        skills = resume_data.get("skills", [])
        if skills:
            doc.add_heading("技能", level=1)
            skill_names = [s.get("skill_name", "") for s in skills]
            doc.add_paragraph(", ".join(skill_names))
        
        # 保存文件
        filename = f"resume_{resume_data.get('id', 'draft')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return f"/resumes/{filename}"
    
    def _generate_html(self, resume_data: Dict, template_id: str) -> str:
        """生成HTML简历"""
        template = self.TEMPLATES.get(template_id, self.TEMPLATES["default"])
        
        header = resume_data.get("header", {})
        
        # 构建HTML
        html = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{header.get('name', '简历')} - JobFit简历</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ 
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px;
        }}
        .header {{ text-align: center; margin-bottom: 30px; }}
        .header h1 {{ font-size: 28px; margin-bottom: 10px; }}
        .header .contact {{ color: #666; font-size: 14px; }}
        .section {{ margin-bottom: 25px; }}
        .section h2 {{ 
            font-size: 18px; 
            border-bottom: 2px solid #4F46E5; 
            padding-bottom: 8px;
            margin-bottom: 15px;
            color: #4F46E5;
        }}
        .item {{ margin-bottom: 15px; }}
        .item-header {{ 
            display: flex; 
            justify-content: space-between; 
            margin-bottom: 5px;
        }}
        .item-title {{ font-weight: bold; }}
        .item-subtitle {{ color: #666; font-size: 14px; }}
        .item-date {{ color: #999; font-size: 14px; }}
        .skills {{ display: flex; flex-wrap: wrap; gap: 8px; }}
        .skill-tag {{ 
            background: #f3f4f6; 
            padding: 4px 12px; 
            border-radius: 15px;
            font-size: 14px;
        }}
        ul {{ margin-left: 20px; }}
        li {{ margin-bottom: 5px; }}
        @media print {{
            body {{ padding: 20px; }}
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{header.get('name', '')}</h1>
        <div class="contact">
            {header.get('phone', '')} | {header.get('email', '')} | {header.get('location', '')}
            {f" | {header.get('linkedin_url', '')}" if header.get('linkedin_url') else ""}
            {f" | {header.get('github_url', '')}" if header.get('github_url') else ""}
        </div>
    </div>
"""
        
        # 个人简介
        if resume_data.get("summary"):
            html += f"""
    <div class="section">
        <h2>个人简介</h2>
        <p>{resume_data['summary']}</p>
    </div>
"""
        
        # 工作经历
        experience = resume_data.get("experience", [])
        if experience:
            html += '<div class="section"><h2>工作经历</h2>'
            for exp in experience:
                html += f"""
        <div class="item">
            <div class="item-header">
                <span class="item-title">{exp.get('position', '')} - {exp.get('company_name', '')}</span>
                <span class="item-date">{exp.get('start_date', '')} - {exp.get('end_date', '至今')}</span>
            </div>
"""
                for ach in exp.get("achievements", []):
                    html += f"<li>{ach}</li>"
                html += "</div>"
            html += "</div>"
        
        # 教育背景
        education = resume_data.get("education", [])
        if education:
            html += '<div class="section"><h2>教育背景</h2>'
            for edu in education:
                html += f"""
        <div class="item">
            <div class="item-header">
                <span class="item-title">{edu.get('school_name', '')} - {edu.get('major', '')}</span>
                <span class="item-date">{edu.get('start_date', '')} - {edu.get('end_date', '')}</span>
            </div>
            <div class="item-subtitle">{edu.get('degree', '')} {f"| GPA: {edu.get('gpa', '')}" if edu.get('gpa') else ""}</div>
        </div>
"""
            html += "</div>"
        
        # 项目经历
        projects = resume_data.get("projects", [])
        if projects:
            html += '<div class="section"><h2>项目经历</h2>'
            for proj in projects:
                html += f"""
        <div class="item">
            <div class="item-header">
                <span class="item-title">{proj.get('project_name', '')}</span>
                <span class="item-date">{proj.get('start_date', '')} - {proj.get('end_date', '')}</span>
            </div>
            {f"<p>{proj.get('description', '')}</p>" if proj.get('description') else ""}
            {f"<p>技术栈: {', '.join(proj.get('tech_stack', []))}</p>" if proj.get('tech_stack') else ""}
        </div>
"""
            html += "</div>"
        
        # 技能
        skills = resume_data.get("skills", [])
        if skills:
            html += """
        <div class="section">
            <h2>技能</h2>
            <div class="skills">
"""
            for skill in skills:
                html += f'<span class="skill-tag">{skill.get("skill_name", "")}</span>'
            html += """
            </div>
        </div>
"""
        
        html += """
    <div style="text-align: center; margin-top: 40px; color: #999; font-size: 12px;">
        由 JobFit AI 简历生成器创建
    </div>
</body>
</html>
"""
        
        return html


# 测试
if __name__ == "__main__":
    exporter = ResumeExporter()
    
    # 测试数据
    test_resume = {
        "id": 1,
        "header": {
            "name": "张三",
            "phone": "138****8888",
            "email": "zhangsan@example.com",
            "location": "北京"
        },
        "summary": "5年后端开发经验，熟练掌握Python、Django、MySQL等技术。",
        "experience": [
            {
                "company_name": "字节跳动",
                "position": "高级后端开发工程师",
                "start_date": "2021.03",
                "end_date": "至今",
                "achievements": ["负责用户增长系统后端开发", "QPS提升50%", "系统可用性99.9%"]
            }
        ],
        "education": [
            {
                "school_name": "清华大学",
                "major": "计算机科学与技术",
                "degree": "本科",
                "start_date": "2013.09",
                "end_date": "2017.06",
                "gpa": "3.8"
            }
        ],
        "skills": [
            {"skill_name": "Python"},
            {"skill_name": "Django"},
            {"skill_name": "MySQL"},
            {"skill_name": "Redis"},
            {"skill_name": "Docker"}
        ]
    }
    
    print("测试生成简历...")
    pdf_url = exporter.export_pdf(test_resume)
    print(f"PDF: {pdf_url}")
    
    word_url = exporter.export_word(test_resume)
    print(f"Word: {word_url}")